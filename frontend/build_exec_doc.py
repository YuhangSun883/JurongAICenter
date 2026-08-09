from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"D:\JRai\JRai_项目阶段性汇报.docx")

NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEXT = "1F2933"
MUTED = "5F6B76"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
BORDER = "D7DEE6"
WHITE = "FFFFFF"
GOLD = "A06A00"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_paragraph_border_bottom(paragraph, color=BORDER, size="8"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.32)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.1

    if "Report Kicker" not in doc.styles:
        style = doc.styles.add_style("Report Kicker", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["Report Kicker"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(GOLD)
    style.paragraph_format.space_after = Pt(5)

    if "Report Subtitle" not in doc.styles:
        style = doc.styles.add_style("Report Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = doc.styles["Report Subtitle"]
    style.font.name = "Calibri"
    style.font.size = Pt(13)
    style.font.color.rgb = RGBColor.from_string(MUTED)
    style.paragraph_format.space_after = Pt(14)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(2)
    r = hp.add_run("JRai  |  项目阶段性汇报")
    set_run_font(r, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(2)
    r = fp.add_run("内部汇报材料  ·  ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(fp, "PAGE")


def add_paragraph(doc, text="", style=None, before=0, after=6, line=1.1, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=TEXT)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=TEXT)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=TEXT)
    return p


def add_section_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_keep_with_next(p)
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 11.5}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(cell, color="BDD2E5", size="10")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120], indent_dxa=120)
    values = [
        ("前端产品框架", "已完成"),
        ("核心工具演示", "9 类入口"),
        ("接口架构", "Mock / Real 双模式"),
    ]
    for cell, (label, value) in zip(table.rows[0].cells, values):
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label)
        set_run_font(r, size=9, color=MUTED, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r = p2.add_run(value)
        set_run_font(r, size=11, color=NAVY, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_status_table(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2200, 4320, 2840], indent_dxa=120)
    headers = ["方向", "当前能力", "业务价值"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell, color="BDD2E5", size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=NAVY, bold=True)

    rows = [
        ("创作平台", "统一首页、导航和多功能工作台", "降低工具切换成本，形成统一入口"),
        ("视频工具", "水印擦除、字幕擦除、画质增强", "覆盖常见视频后期处理场景"),
        ("内容生产", "AI 视频、AI 图片、商品详情套图、Agent", "支持电商素材的多类型生产"),
        ("素材管理", "全局上传和素材选择流程", "为团队资产沉淀和复用打基础"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            set_cell_border(cell, color=BORDER, size="8")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, size=9.5, color=TEXT, bold=(cell is cells[0]))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_two_column_table(doc, left_title, left_items, right_title, right_items):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4680, 4680], indent_dxa=120)
    for cell, title, items, fill in zip(
        table.rows[0].cells,
        [left_title, right_title],
        [left_items, right_items],
        [LIGHT_BLUE, LIGHT_GRAY],
    ):
        set_cell_shading(cell, fill)
        set_cell_border(cell, color=BORDER, size="8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(title)
        set_run_font(r, size=10.5, color=NAVY if fill == LIGHT_BLUE else DARK_BLUE, bold=True)
        for item in items:
            p = cell.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.12)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(item)
            set_run_font(r, size=9.5, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    core = doc.core_properties
    core.title = "JRai 智能内容创作平台 - 项目阶段性汇报"
    core.subject = "面向老板的项目阶段性汇报"
    core.author = ""
    core.keywords = "JRai, AI, 电商, 内容生产, 项目汇报"

    kicker = doc.add_paragraph(style="Report Kicker")
    kicker.paragraph_format.space_before = Pt(4)
    r = kicker.add_run("项目阶段性汇报")
    set_run_font(r, size=10, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    r = title.add_run("JRai 智能内容创作平台")
    set_run_font(r, size=25, color=NAVY, bold=True)

    subtitle = doc.add_paragraph(style="Report Subtitle")
    r = subtitle.add_run("面向电商与短视频团队的一站式 AI 内容生产工作台")
    set_run_font(r, size=13, color=MUTED)

    metadata = [
        ("当前阶段", "已完成前端产品原型和核心交互演示，进入后端接口联调阶段"),
        ("汇报日期", "2026 年 8 月 5 日"),
        ("交付范围", "前端页面、工作台交互、素材流程、接口预留"),
        ("当前环境", "本地可运行，支持 Mock 模式演示"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + "：")
        set_run_font(r, size=10.5, color=TEXT, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=TEXT)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(12)
    set_paragraph_border_bottom(rule, color="B8C9D8", size="10")

    add_callout(
        doc,
        "核心结论",
        "JRai 已从单点页面开发进入平台化前端产品阶段，具备多工具统一入口、素材复用和任务工作流基础，下一步重点是接入真实后端能力。",
    )
    add_metric_strip(doc)

    add_section_heading(doc, "一、项目定位", 1)
    add_paragraph(
        doc,
        "JRai 是一套面向电商和短视频团队的 AI 内容生产工作台，将视频生成、图片生成、商品详情制作、素材管理和视频后期处理集中在一个平台中，帮助团队降低内容制作门槛，提高创作效率。",
    )
    add_paragraph(doc, "平台未来重点服务以下场景：", after=4)
    for item in [
        "电商商品详情页和营销素材制作",
        "短视频批量生产与广告素材生成",
        "视频字幕、水印和画质处理",
        "企业内部内容资产统一管理",
    ]:
        add_bullet(doc, item)

    add_section_heading(doc, "二、项目价值", 1)
    add_paragraph(
        doc,
        "传统内容制作往往需要在多个软件和平台之间切换，流程分散、学习成本高、协作效率低。JRai 的核心价值，是把素材上传、AI 创作、内容处理、任务管理和结果预览整合成一条连续工作流。",
    )
    for item in [
        "统一入口：多个 AI 能力进入同一套平台和导航体系。",
        "流程清晰：围绕上传、配置、提交、预览和任务管理设计。",
        "能力可扩展：新增工具可以复用现有工作台和素材流程。",
        "适合产品化：后续可以扩展账号、积分、套餐、团队空间和 API 服务。",
    ]:
        add_bullet(doc, item)

    add_section_heading(doc, "三、当前成果", 1)
    add_status_table(doc)

    add_section_heading(doc, "四、项目核心卖点", 1)
    add_section_heading(doc, "1. 一站式 AI 内容生产", 2)
    add_paragraph(doc, "JRai 不只是单一 AI 工具，而是围绕电商内容生产建立完整工作台，覆盖从素材到成片的多个环节。")
    add_section_heading(doc, "2. 面向实际业务流程", 2)
    for item in [
        "上传素材",
        "选择创作工具",
        "配置生成参数",
        "提交处理任务",
        "查看预览结果",
        "管理历史任务",
    ]:
        add_number(doc, item)
    add_section_heading(doc, "3. 工具模块可持续扩展", 2)
    add_paragraph(doc, "水印擦除、字幕擦除、画质增强、商品详情套图等功能均采用独立页面和独立工作台设计，后续可以快速扩展数字人、模特换装、爆款裂变等能力。")
    add_section_heading(doc, "4. 前后端解耦，降低试错成本", 2)
    add_paragraph(doc, "前端可以独立完成产品流程和用户体验验证，后端可以并行开发。产品方可以先确认页面和交互，再逐步接入真实 AI 服务。")

    add_section_heading(doc, "五、当前状态与边界", 1)
    add_two_column_table(
        doc,
        "已完成",
        [
            "前端主体框架与统一导航",
            "核心创作页面和视频工具工作台",
            "全局素材上传和选择流程",
            "Mock / Real 双模式接口架构",
            "本地演示环境和核心页面验证",
        ],
        "待接入",
        [
            "真实登录、注册和权限",
            "真实视频上传和 AI 处理任务",
            "任务进度、失败重试和结果下载",
            "积分、套餐、支付和历史作品",
            "团队协作和企业级权限管理",
        ],
    )
    add_callout(
        doc,
        "说明",
        "当前版本是前端可演示版本，字幕擦除和画质增强中的任务队列主要用于展示交互流程，真实处理结果需要等待后端接口接入。",
    )

    add_section_heading(doc, "六、下一步计划", 1)
    add_section_heading(doc, "第一阶段：接入后端能力", 2)
    add_paragraph(doc, "优先完成登录注册、素材上传、字幕擦除、水印擦除、画质增强和任务进度接口，形成可真实闭环的演示版本。")
    add_section_heading(doc, "第二阶段：完善产品闭环", 2)
    add_paragraph(doc, "补齐积分套餐、历史作品、失败重试、成品下载、素材搜索和团队协作能力。")
    add_section_heading(doc, "第三阶段：扩展商业化能力", 2)
    add_paragraph(doc, "围绕企业账号、团队空间、会员套餐、批量生成、行业模板和品牌资产库进行产品化扩展。")

    add_section_heading(doc, "七、重点演示入口", 1)
    for item in [
        "首页：/",
        "商品详情套图：/tools/product-image",
        "字幕擦除：/tools/subtitle-remover",
        "画质增强：/tools/image-enhancer",
        "水印擦除：/tools/watermark-remover",
    ]:
        add_bullet(doc, item)
    add_paragraph(doc, "本地启动后访问：http://localhost:3000", before=2, after=8)

    add_section_heading(doc, "八、总结", 1)
    add_paragraph(
        doc,
        "JRai 已经具备一个 AI 内容生产平台的前端雏形：有统一入口、有核心工具、有素材流程、有任务工作流，也为登录、积分、任务和商业化能力预留了扩展空间。接入后端后，项目可以快速从前端演示版本升级为面向电商和短视频团队的实际生产平台。",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
